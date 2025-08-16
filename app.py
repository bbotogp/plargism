from flask import Flask, request, render_template_string, send_file
import requests, time, re
from lxml import etree
from datetime import datetime, date, timedelta
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document

app = Flask(__name__)

INTRO_TITLES = {"introduction", "background", "intro", "introduction/background"}
HEADERS = {"User-Agent": "MedIntroBot/1.0 (mailto:example@example.com)"}

def build_query(q, days=7):
    d1 = date.today() - timedelta(days=days)
    d2 = date.today()
    base_q = f'OPEN_ACCESS:y AND HAS_FT:y AND PUB_TYPE:"Journal Article" AND FIRST_PDATE:[{d1} TO {d2}]'
    return f'({base_q}) AND ({q})' if q else base_q

def search_epmc(query, max_results=5):
    url = "https://www.ebi.ac.uk/europepmc/webservices/rest/search"
    items = []
    page = 1
    while len(items) < max_results:
        params = {
            "query": query, "resultType": "core", "format": "json",
            "pageSize": min(100, max_results - len(items)),
            "page": page, "sort": "P_PDATE_D"
        }
        r = requests.get(url, params=params, headers=HEADERS)
        data = r.json()
        hits = data.get("resultList", {}).get("result", [])
        if not hits: break
        items.extend(hits)
        page += 1
    return items[:max_results]

def get_intro_from_pmcid(pmcid):
    url = f"https://www.ebi.ac.uk/europepmc/webservices/rest/{pmcid}/fullTextXML"
    r = requests.get(url, headers=HEADERS)
    if r.status_code != 200:
        return None
    parser = etree.XMLParser(recover=True)
    root = etree.fromstring(r.content, parser=parser)
    bodies = root.findall(".//body")
    if not bodies:
        return None
    body = bodies[0]
    for sec in body.findall(".//sec"):
        title_elem = sec.find("title")
        if title_elem is not None:
            title_text = "".join(title_elem.itertext()).strip().lower()
            if title_text in INTRO_TITLES:
                return extract_text(sec)
    first_sec = body.find(".//sec")
    return extract_text(first_sec) if first_sec is not None else None

def extract_text(elem):
    text_parts = []
    for p in elem.findall(".//p"):
        txt = "".join(p.itertext()).strip()
        if txt:
            text_parts.append(txt)
    return "\n\n".join(text_parts)

def build_result(entry, intro_text):
    authors = entry.get("authorString", "").split(", ")
    return {
        "title": entry.get("title"),
        "authors": authors,
        "journal": entry.get("journalTitle"),
        "year": entry.get("pubYear"),
        "doi": entry.get("doi"),
        "pmcid": entry.get("pmcid"),
        "url": f"https://europepmc.org/article/pmc/{entry.get('pmcid')}",
        "license": entry.get("license"),
        "intro_text": intro_text
    }

def export_pdf(results, filename="output.pdf"):
    doc = SimpleDocTemplate(filename, pagesize=A4)
    styles = getSampleStyleSheet()
    flow = [Paragraph("نتائج البحث", styles["Title"]), Spacer(1, 12)]
    for r in results:
        flow.append(Paragraph(f"{r['title']}", styles["Heading2"]))
        flow.append(Paragraph(f"الكاتب: {', '.join(r['authors'])}", styles["Normal"]))
        flow.append(Paragraph(f"المجلة: {r['journal']} ({r['year']})", styles["Normal"]))
        flow.append(Paragraph(f"الرابط: {r['url']}", styles["Normal"]))
        if r["intro_text"]:
            flow.append(Paragraph("📘 المقدمة:", styles["Normal"]))
            flow.append(Paragraph(r["intro_text"], styles["Normal"]))
        else:
            flow.append(Paragraph("⚠️ المقدمة غير متوفرة.", styles["Normal"]))
        flow.append(Spacer(1, 16))
    doc.build(flow)
    return filename

def export_docx(results, filename="output.docx"):
    doc = Document()
    doc.add_heading("نتائج البحث", 0)
    for r in results:
        doc.add_heading(r["title"], level=2)
        doc.add_paragraph(f"الكاتب: {', '.join(r['authors'])}")
        doc.add_paragraph(f"المجلة: {r['journal']} ({r['year']})")
        doc.add_paragraph(f"الرابط: {r['url']}")
        doc.add_paragraph("المقدمة:" if r["intro_text"] else "⚠️ المقدمة غير متوفرة")
        if r["intro_text"]:
            doc.add_paragraph(r["intro_text"])
    doc.save(filename)
    return filename

HTML_TEMPLATE = """<!doctype html><html lang="ar" dir="rtl"><head>
  <meta charset="utf-8"><title>مقدمات الأبحاث الطبية</title>
  <style>body{font-family:sans-serif;direction:rtl;background:#f5f5f5;padding:20px}
  .box{background:white;padding:20px;max-width:800px;margin:auto;border-radius:10px;box-shadow:0 0 10px #ccc}
  input,button{padding:10px;width:100%;margin-top:10px} .res{margin-top:30px} .intro{background:#eef;padding:10px;border-radius:5px;white-space:pre-line}</style>
</head><body><div class="box">
  <h1>📄 استخراج مقدمة الأبحاث الطبية</h1>
  <form method="post">
    <input name="query" placeholder="مثال: covid-19" required value="{{query}}">
    <input name="days" type="number" min="1" max="90" value="{{days}}">
    <input name="max" type="number" min="1" max="20" value="{{max}}">
    <button name="export" value="search">بحث</button>
    <button name="export" value="pdf">📥 تحميل PDF</button>
    <button name="export" value="docx">📄 تحميل Word</button>
  </form>
  {% if results %}
    <div class="res">
      {% for r in results %}
        <h2>{{loop.index}}. {{r.title}}</h2>
        <p><b>المؤلفون:</b> {{r.authors|join(", ")}}<br>
        <b>المجلة:</b> {{r.journal}} ({{r.year}})<br>
        <b><a href="{{r.url}}" target="_blank">رابط المقال</a></b></p>
        {% if r.intro_text %}
        <div class="intro">{{r.intro_text}}</div>
        {% else %}<p style="color:red;">⚠️ لا توجد مقدمة متاحة</p>{% endif %}
        <hr>
      {% endfor %}
    </div>
  {% endif %}
</div></body></html>
"""

@app.route("/", methods=["GET", "POST"])
def index():
    query = request.form.get("query", "")
    days = int(request.form.get("days", 7))
    max_results = int(request.form.get("max", 5))
    export = request.form.get("export", "search")
    results = []

    if request.method == "POST":
        q = build_query(query, days)
        entries = search_epmc(q, max_results)
        for e in entries:
            pmcid = e.get("pmcid")
            intro = get_intro_from_pmcid(pmcid) if pmcid else None
            results.append(build_result(e, intro))
            time.sleep(0.2)

        if export == "pdf":
            f = export_pdf(results)
            return send_file(f, as_attachment=True)
        elif export == "docx":
            f = export_docx(results)
            return send_file(f, as_attachment=True)

    return render_template_string(HTML_TEMPLATE, results=results, query=query, days=days, max=max_results)

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=7860)
